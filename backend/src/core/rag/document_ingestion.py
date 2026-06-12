import os
import re
import tempfile
import threading
from io import BytesIO
from pathlib import Path
from typing import Any

from langchain_core.documents import Document

from src.core.rag.ingest import split_text
from src.core.rag.qdrant_store import QdrantKnowledgeStore

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}
OCR_ENABLED = os.getenv("RAG_OCR_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
OCR_LANGUAGE = os.getenv("RAG_OCR_LANGUAGE", "rus+eng")
OCR_MIN_CHARS_PER_PAGE = int(os.getenv("RAG_OCR_MIN_CHARS_PER_PAGE", "40"))
OCR_CONCURRENCY = max(1, int(os.getenv("RAG_OCR_CONCURRENCY", "1")))
_OCR_SEMAPHORE = threading.Semaphore(OCR_CONCURRENCY)


class DocumentExtractionError(RuntimeError):
    pass


def index_knowledge_file(
    prepared_file: dict[str, Any],
    theme_id: str,
    theme_title: str,
    store: QdrantKnowledgeStore,
) -> int:
    documents = extract_documents(
        data=prepared_file["data"],
        filename=prepared_file["filename"],
        content_type=prepared_file["content_type"],
        file_id=str(prepared_file["file_id"]),
        theme_id=theme_id,
        theme_title=theme_title,
    )
    documents = prepare_documents_for_chunking(documents, fallback_section=theme_title)
    chunks = split_text(documents)
    if not chunks:
        raise DocumentExtractionError("No chunks were produced from the uploaded file")
    return store.upsert_documents(chunks, incremental=True)


def extract_documents(
    data: bytes,
    filename: str,
    content_type: str,
    file_id: str,
    theme_id: str,
    theme_title: str,
) -> list[Document]:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise DocumentExtractionError(f"Unsupported file extension: {ext or '<empty>'}")

    metadata = _base_metadata(
        filename=filename,
        content_type=content_type,
        file_id=file_id,
        theme_id=theme_id,
        theme_title=theme_title,
    )
    if ext == ".txt":
        documents = _extract_txt(data, metadata)
    elif ext == ".docx":
        documents = _extract_docx(data, metadata)
    else:
        documents = _extract_pdf(data, filename, metadata)

    documents = [doc for doc in documents if doc.page_content.strip()]
    if not documents:
        raise DocumentExtractionError("No text could be extracted from the uploaded file")
    return documents


def prepare_documents_for_chunking(
    documents: list[Document],
    fallback_section: str | None = None,
) -> list[Document]:
    prepared: list[Document] = []
    current_section = fallback_section or None
    section_index = -1
    source_order = 0

    for document in documents:
        section_documents, current_section, section_index = _split_document_by_sections(
            document=document,
            current_section=current_section,
            section_index=section_index,
            fallback_section=fallback_section,
        )
        for section_document in section_documents:
            section_document.metadata["document_order"] = source_order
            source_order += 1
            prepared.append(section_document)

    return prepared


def _base_metadata(
    filename: str,
    content_type: str,
    file_id: str,
    theme_id: str,
    theme_title: str,
) -> dict[str, Any]:
    return {
        "source": f"admin-upload:{file_id}:{filename}",
        "source_type": "admin_upload",
        "source_theme": theme_title,
        "theme_id": theme_id,
        "file_id": file_id,
        "filename": filename,
        "content_type": content_type,
    }


def _extract_txt(data: bytes, metadata: dict[str, Any]) -> list[Document]:
    text = _decode_text(data)
    return [
        Document(
            page_content=_normalize_text(text),
            metadata={**metadata, "page": 1, "extractor": "text"},
        )
    ]


def _extract_docx(data: bytes, metadata: dict[str, Any]) -> list[Document]:
    try:
        from docx import Document as DocxDocument
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise DocumentExtractionError("python-docx is not installed") from exc

    docx = DocxDocument(BytesIO(data))
    parts: list[str] = []
    for child in docx.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, docx)
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style_name.startswith("heading"):
                level = _heading_level(style_name)
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
        elif isinstance(child, CT_Tbl):
            table_md = _table_to_markdown(Table(child, docx))
            if table_md:
                parts.append(table_md)

    text = _normalize_text("\n\n".join(parts))
    return [
        Document(
            page_content=text,
            metadata={**metadata, "page": 1, "extractor": "python-docx"},
        )
    ]


def _extract_pdf(data: bytes, filename: str, metadata: dict[str, Any]) -> list[Document]:
    if OCR_ENABLED:
        with _OCR_SEMAPHORE:
            return _extract_pdf_inner(data, filename, metadata)
    return _extract_pdf_inner(data, filename, metadata)


def _extract_pdf_inner(data: bytes, filename: str, metadata: dict[str, Any]) -> list[Document]:
    try:
        import pymupdf4llm
    except ImportError as exc:
        raise DocumentExtractionError("pymupdf4llm is not installed") from exc

    suffix = Path(filename).suffix or ".pdf"
    with tempfile.NamedTemporaryFile(suffix=suffix) as tmp:
        tmp.write(data)
        tmp.flush()
        markdown_kwargs = {
            "page_chunks": True,
            "use_ocr": OCR_ENABLED,
            "ocr_language": OCR_LANGUAGE,
        }
        try:
            page_chunks = pymupdf4llm.to_markdown(tmp.name, **markdown_kwargs)
        except TypeError:
            page_chunks = pymupdf4llm.to_markdown(tmp.name, page_chunks=True)
        documents = _pdf_chunks_to_documents(page_chunks, metadata)
        if OCR_ENABLED:
            documents = _replace_blank_pdf_pages_with_ocr(tmp.name, documents, metadata)
    return documents


def _pdf_chunks_to_documents(page_chunks: list[dict[str, Any]], metadata: dict[str, Any]) -> list[Document]:
    documents: list[Document] = []
    for index, chunk in enumerate(page_chunks):
        chunk_metadata = chunk.get("metadata") or {}
        page_number = chunk_metadata.get("page_number") or chunk_metadata.get("page") or index + 1
        text = _normalize_text(str(chunk.get("text") or ""))
        documents.append(
            Document(
                page_content=text,
                metadata={
                    **metadata,
                    "page": int(page_number),
                    "extractor": "pymupdf4llm",
                    "pdf_title": chunk_metadata.get("title"),
                },
            )
        )
    return documents


def _replace_blank_pdf_pages_with_ocr(
    pdf_path: str,
    documents: list[Document],
    metadata: dict[str, Any],
) -> list[Document]:
    blank_pages = [
        index
        for index, document in enumerate(documents)
        if len(_semantic_text(document.page_content)) < OCR_MIN_CHARS_PER_PAGE
    ]
    if not blank_pages:
        return documents

    ocr_documents = _extract_pdf_pages_with_ocr(pdf_path, blank_pages, metadata)
    if not ocr_documents:
        return documents

    by_page = {document.metadata["page"]: document for document in ocr_documents}
    replaced: list[Document] = []
    for document in documents:
        page_number = document.metadata.get("page")
        replacement = by_page.get(page_number)
        replaced.append(replacement if replacement and replacement.page_content.strip() else document)
    return replaced


def _extract_pdf_pages_with_ocr(
    pdf_path: str,
    page_indexes: list[int],
    metadata: dict[str, Any],
) -> list[Document]:
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError:
        return []

    documents: list[Document] = []
    with fitz.open(pdf_path) as pdf:
        for page_index in page_indexes:
            page = pdf[page_index]
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.open(BytesIO(pixmap.tobytes("png")))
            text = pytesseract.image_to_string(image, lang=OCR_LANGUAGE)
            documents.append(
                Document(
                    page_content=_normalize_text(text),
                    metadata={
                        **metadata,
                        "page": page_index + 1,
                        "extractor": "tesseract-ocr",
                        "ocr_language": OCR_LANGUAGE,
                    },
                )
            )
    return documents


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def _split_document_by_sections(
    document: Document,
    current_section: str | None,
    section_index: int,
    fallback_section: str | None,
) -> tuple[list[Document], str | None, int]:
    sections: list[Document] = []
    buffer: list[str] = []
    buffer_section = current_section or fallback_section
    buffer_section_index = max(section_index, 0)

    def flush() -> None:
        text = _normalize_text("\n".join(buffer))
        if not text:
            return
        sections.append(
            Document(
                page_content=text,
                metadata={
                    **document.metadata,
                    "section_title": buffer_section or fallback_section,
                    "section_index": buffer_section_index,
                    "page_start": document.metadata.get("page"),
                    "page_end": document.metadata.get("page"),
                    "document_chars": len(text),
                },
            )
        )

    for line in document.page_content.splitlines():
        heading = _section_heading(line)
        if heading:
            flush()
            buffer = [line]
            current_section = heading
            section_index += 1
            buffer_section = current_section
            buffer_section_index = section_index
            continue
        buffer.append(line)

    flush()
    return sections, current_section, section_index


def _section_heading(line: str) -> str | None:
    text = line.strip()
    if not text:
        return None

    markdown_match = re.match(r"^#{1,6}\s+(.+)$", text)
    if markdown_match:
        return _clean_heading(markdown_match.group(1))

    if not _looks_like_plain_heading(text):
        return None
    return _clean_heading(text)


def _looks_like_plain_heading(text: str) -> bool:
    if len(text) > 160 or "|" in text:
        return False
    if text.endswith((".", ",", ";", ":")):
        return False
    if len(text.split()) > 14:
        return False

    if re.match(r"^(?:\d+(?:\.\d+)+\.?\s+|[IVXLC]+[\).\s]+)", text, re.IGNORECASE):
        return True
    if re.match(r"^(?:раздел|глава|блок|тема|приложение)\b", text, re.IGNORECASE):
        return True
    return False


def _clean_heading(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip("#* \t")


def _semantic_text(text: str) -> str:
    return re.sub(r"\s+", "", text)


def _heading_level(style_name: str) -> int:
    match = re.search(r"(\d+)", style_name)
    if not match:
        return 2
    return max(1, min(int(match.group(1)), 6))


def _table_to_markdown(table: Any) -> str:
    rows = [[_escape_markdown_cell(cell.text.strip()) for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    separator = ["---"] * width
    body = normalized_rows[1:]
    md_rows = [header, separator, *body]
    return "\n".join("| " + " | ".join(row) + " |" for row in md_rows)


def _escape_markdown_cell(text: str) -> str:
    return re.sub(r"\s+", " ", text).replace("|", "\\|")
