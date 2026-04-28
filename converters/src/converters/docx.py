from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.converters.base import BaseToMarkdownConverter


class DocxToMarkdownConverter(BaseToMarkdownConverter):
    file_extension_length = 5  # Учитываем точку: *.docx

    def _get_doc_elements(self, doc_name: str) -> list[Paragraph | Table]:
        doc = Document(f'{self._docs_dir}/{doc_name}.docx')
        elements = []

        for el in doc.element.body:
            if el.tag.endswith("p"):
                for paragraph in doc.paragraphs:
                    if paragraph._element == el:
                        elements.append(paragraph)
                        break

            elif el.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == el:
                        elements.append(table)
                        break

        return elements

    def _sync_convert_document_to_markdown(self, doc_name: str) -> None:
        element_list = self._get_doc_elements(doc_name)
        md_lines = []

        previous_header_idx = 0
        previous_header_is_specific = False
        paragraph_idx = 0

        for element in element_list:
            if type(element) == Paragraph:
                paragraph = element
                text = paragraph.text

                if not text:
                    continue

                if self._is_header(paragraph):
                    idx_diff = abs(paragraph_idx - previous_header_idx)
                    if idx_diff == 2:
                        if (self._is_specific_header(paragraph) and previous_header_is_specific) or (
                                not self._is_specific_header(paragraph) and not previous_header_is_specific):
                            md_lines[previous_header_idx] = f"{md_lines[previous_header_idx]} {text}"
                            continue
                    if self._is_specific_header(paragraph):
                        text = f"# {text}"
                        previous_header_is_specific = True
                    else:
                        text = f"## {text}"
                        previous_header_is_specific = False
                    previous_header_idx = paragraph_idx

                text = self._convert_numbers(text)

                md_lines.append(text)
                md_lines.append("")
                paragraph_idx += 2  # Увеличиваем на 2 из-за пустой строки

            if type(element) == Table:
                md_lines.append(self._process_table(element))
                md_lines.append("")
                paragraph_idx += 2

        self._write_in_md_file(md_lines, doc_name)

    @staticmethod
    def _process_table(table: Table) -> str:
        markdown_table = []

        headers = [cell.text.strip().replace("\n", "") for cell in table.rows[0].cells]
        markdown_table.append(f"| {' | '.join(headers)} |")

        markdown_table.append(f"|{'|'.join([' --- '] * len(headers))}|")

        for row in table.rows[1:]:
            row_data = [cell.text.strip().replace("\n", "") for cell in row.cells]
            markdown_table.append(f"| {' | '.join(row_data)} |")

        return '\n'.join(markdown_table)

    @staticmethod
    def _is_header(string: Paragraph) -> bool:
        if (
                string.style.name == "ConsPlusTitle"
                or all(run.font.bold for run in string.runs)
                or string.alignment == WD_ALIGN_PARAGRAPH.CENTER
        ):
            return True
        return False

    @staticmethod
    def _is_specific_header(paragraph: Paragraph) -> bool:
        specific_header_list = [
            "ФЕДЕРАЛЬНАЯ СЛУЖБА ПО НАДЗОРУ В СФЕРЕ ЗАЩИТЫ",
            "ПРАВ ПОТРЕБИТЕЛЕЙ И БЛАГОПОЛУЧИЯ ЧЕЛОВЕКА",
            "ГЛАВНЫЙ ГОСУДАРСТВЕННЫЙ САНИТАРНЫЙ ВРАЧ"
            "РОССИЙСКОЙ ФЕДЕРАЦИИ",
            "ПОСТАНОВЛЕНИЕ",
            "от 28 января 2021 г. N 4",
            "от 28 января 2021 г. N 3",
            "ОБ УТВЕРЖДЕНИИ САНИТАРНЫХ ПРАВИЛ И НОРМ САНПИН 3.3686-21",
            '"САНИТАРНО-ЭПИДЕМИОЛОГИЧЕСКИЕ ТРЕБОВАНИЯ ПО ПРОФИЛАКТИКЕ',
            'ИНФЕКЦИОННЫХ БОЛЕЗНЕЙ"',
            "ГЛАВНЫЙ ГОСУДАРСТВЕННЫЙ САНИТАРНЫЙ ВРАЧ",
            "РОССИЙСКОЙ ФЕДЕРАЦИИ",
        ]

        if paragraph.text in specific_header_list:
            return True

        return False
